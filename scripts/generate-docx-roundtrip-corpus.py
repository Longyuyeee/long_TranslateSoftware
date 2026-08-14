"""Generate deterministic, synthetic DOCX fixtures for rebuild compatibility tests."""

from __future__ import annotations

import argparse
import base64
import io
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


FIXED_TIME = datetime(2026, 8, 14, tzinfo=timezone.utc)
ZIP_TIME = (2026, 8, 14, 0, 0, 0)


def set_properties(document: Document, title: str) -> None:
    document.core_properties.title = title
    document.core_properties.author = "Long Translate synthetic corpus"
    document.core_properties.created = FIXED_TIME
    document.core_properties.modified = FIXED_TIME


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    properties.append(color)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction: str, display: str) -> None:
    run = paragraph.add_run()._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, code, separate, text, end):
        run.append(node)


def mark_numbered(paragraph, number_id: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    identifier = OxmlElement("w:numId")
    identifier.set(qn("w:val"), str(number_id))
    numbering.extend((level, identifier))
    properties.append(numbering)


def tiny_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAYAAABytg0k"
        "AAAAFElEQVR4nGNUTX79n4GBgYGJAQoAJJICdqnjmg8AAAAASUVORK5CYII="
    )


def normalize_docx(source: Path) -> None:
    entries: list[tuple[str, bytes, int]] = []
    with zipfile.ZipFile(source, "r") as archive:
        for info in archive.infolist():
            entries.append((info.filename, archive.read(info), info.compress_type))
    temporary = source.with_suffix(".normalized.tmp")
    with zipfile.ZipFile(temporary, "w") as archive:
        for name, payload, compression in sorted(entries):
            info = zipfile.ZipInfo(name, ZIP_TIME)
            info.compress_type = compression
            info.external_attr = 0o600 << 16
            archive.writestr(info, payload)
    temporary.replace(source)


def save(document: Document, output: Path, title: str) -> None:
    set_properties(document, title)
    document.save(output)
    normalize_docx(output)


def create_corpus(output: Path) -> None:
    output.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading("Product overview", level=1)
    paragraph = document.add_paragraph("Long Translate keeps formatting and follows ")
    add_hyperlink(paragraph, "the project guide", "https://example.com/guide")
    paragraph.add_run(" safely.")
    save(document, output / "01-heading-hyperlink.docx", "Heading and hyperlink")

    document = Document()
    document.add_heading("Release checklist", level=2)
    for value in ("Inspect the source", "Translate each segment", "Verify the output"):
        paragraph = document.add_paragraph(value, style="List Bullet")
        mark_numbered(paragraph, 1)
    for value in ("First gate", "Second gate"):
        paragraph = document.add_paragraph(value, style="List Number")
        mark_numbered(paragraph, 5)
    save(document, output / "02-lists.docx", "Lists")

    document = Document()
    document.add_heading("Compatibility matrix", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Application"
    table.cell(0, 1).text = "Result"
    table.cell(1, 0).text = "LibreOffice"
    nested = table.cell(1, 1).add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "Open"
    nested.cell(0, 1).text = "Pass"
    save(document, output / "03-tables.docx", "Tables")

    document = Document()
    document.sections[0].header.paragraphs[0].text = "Primary header"
    document.sections[0].footer.paragraphs[0].text = "Primary footer"
    document.add_paragraph("First section body")
    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.header.is_linked_to_previous = False
    second.footer.is_linked_to_previous = False
    second.header.paragraphs[0].text = "Secondary header"
    second.footer.paragraphs[0].text = "Secondary footer"
    document.add_paragraph("Second section body")
    save(document, output / "04-sections.docx", "Sections")

    document = Document()
    document.add_heading("Unicode and resources", level=1)
    document.add_paragraph("中文 日本語 한국어 café 👩‍💻 e\u0301")
    image = io.BytesIO(tiny_png())
    document.add_picture(image, width=Inches(0.2))
    field = document.add_paragraph("Generated field: ")
    add_field(field, "DATE", "2026-08-14")
    document.add_paragraph("Second page").add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("Content after the page break")
    save(document, output / "05-unicode-resources.docx", "Unicode and resources")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "output",
        nargs="?",
        type=Path,
        default=Path("src-tauri/tests/fixtures/docx/roundtrip"),
    )
    args = parser.parse_args()
    create_corpus(args.output)


if __name__ == "__main__":
    main()
